import fitz  # PyMuPDF - fastest PDF parser
from docx import Document
import httpx
import asyncio
from concurrent.futures import ThreadPoolExecutor
import email
from email.mime.text import MIMEText
from typing import Union, Optional, Dict
import os
import tempfile
import hashlib
from urllib.parse import urlparse
from functools import lru_cache
import aiofiles
import time

# Global caches for performance
parsed_cache: Dict[str, str] = {}
file_hash_cache: Dict[str, str] = {}

# Thread pool for CPU-intensive tasks
thread_pool = ThreadPoolExecutor(max_workers=4)

class DocumentParser:
    """High-performance async document parser with caching"""
    
    def __init__(self):
        self.client = httpx.AsyncClient(timeout=30.0)
    
    async def __aenter__(self):
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.client.aclose()
    
    @staticmethod
    def get_content_hash(content: Union[str, bytes]) -> str:
        """Generate hash for content-based caching"""
        if isinstance(content, str):
            content = content.encode('utf-8')
        return hashlib.sha256(content).hexdigest()[:16]
    
    async def download_file_async(self, url: str) -> str:
        """Async download with caching"""
        try:
            # Check cache first
            url_hash = self.get_content_hash(url)
            if url_hash in file_hash_cache:
                cached_path = file_hash_cache[url_hash]
                if os.path.exists(cached_path):
                    return cached_path
            
            response = await self.client.get(url)
            response.raise_for_status()
            
            # Get file extension from URL or content-type
            parsed_url = urlparse(url)
            file_extension = os.path.splitext(parsed_url.path)[1]
            
            if not file_extension:
                content_type = response.headers.get('content-type', '')
                if 'pdf' in content_type:
                    file_extension = '.pdf'
                elif 'word' in content_type or 'docx' in content_type:
                    file_extension = '.docx'
                else:
                    file_extension = '.txt'
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                temp_file.write(response.content)
                return temp_file.name
                
        except Exception as e:
            raise Exception(f"Failed to download file: {str(e)}")
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
                text += "\n\n"  # Add page break
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_pdf_sync(file_path: str) -> str:
        """Synchronous PDF parsing method for compatibility"""
        return DocumentParser.parse_pdf(file_path)
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_email(content: str) -> str:
        """Parse email content"""
        try:
            # If it's already plain text, return as is
            if not content.startswith('From:') and not content.startswith('Subject:'):
                return content
            
            # Try to parse as email
            msg = email.message_from_string(content)
            
            # Extract subject and body
            subject = msg.get('Subject', '')
            body = ""
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True).decode('utf-8')
                        break
            else:
                body = msg.get_payload(decode=True).decode('utf-8')
            
            return f"Subject: {subject}\n\n{body}".strip()
            
        except Exception as e:
            # If email parsing fails, return content as is
            return content
    
    @classmethod
    def parse_document(cls, url_or_content: str) -> str:
        """Main method to parse document from URL or direct content"""
        try:
            # Check if it's a URL
            if url_or_content.startswith(('http://', 'https://')):
                # Download file
                file_path = cls.download_file(url_or_content)
                
                try:
                    # Determine file type and parse
                    if file_path.lower().endswith('.pdf'):
                        text = cls.parse_pdf(file_path)
                    elif file_path.lower().endswith(('.docx', '.doc')):
                        text = cls.parse_docx(file_path)
                    else:
                        # Try to read as text file
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                finally:
                    # Clean up temporary file
                    if os.path.exists(file_path):
                        os.unlink(file_path)
                
                return text
            
            else:
                # Treat as direct content (email or text)
                return cls.parse_email(url_or_content)
                
        except Exception as e:
            raise Exception(f"Document parsing failed: {str(e)}")

# Convenience function
def parse_document(url_or_content: str) -> str:
    """Parse document and return extracted text"""
    parser = DocumentParser()
    return parser.parse_document(url_or_content)
